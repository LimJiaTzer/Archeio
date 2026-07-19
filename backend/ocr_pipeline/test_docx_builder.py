import io
import unittest

import cv2
import numpy as np

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ocr_pipeline.docx_builder import build_docx
from ocr_pipeline import layout
from ocr_pipeline.layout import Region
from ocr_pipeline.pipeline import _classify_region, _crop_figure, _infer_alignment, _infer_heading_level
from ocr_pipeline.style import TextStyle


class DocxBuilderTests(unittest.TestCase):
    def test_classifies_layout_regions_and_alignment(self):
        self.assertEqual(_classify_region(Region(kind="title", bbox=[0, 0, 1, 1])), "heading")
        self.assertEqual(_classify_region(Region(kind="text", bbox=[0, 0, 1, 1])), "paragraph")
        self.assertEqual(_classify_region(Region(kind="list", bbox=[0, 0, 1, 1])), "list")
        self.assertEqual(_classify_region(Region(kind="image", bbox=[0, 0, 1, 1])), "figure")
        self.assertEqual(_infer_alignment([420, 80, 1180, 160], 1600), "center")
        self.assertEqual(_infer_alignment([120, 180, 1320, 230], 1600), "left")
        # Full-width region with narrow, near-symmetric margins on both
        # sides -- conventionally justified body text (verified against
        # PaddleOCR's own native docx export on an equivalent block).
        self.assertEqual(_infer_alignment([15, 46, 355, 541], 370), "justify")

    def test_preserves_region_semantics_and_text_style(self):
        region = Region(
            kind="title",
            bbox=[300, 80, 1300, 160],
            role="heading",
            text="Centered heading",
            alignment="center",
            style=TextStyle(
                font_size_pt=18,
                color_rgb=(10, 20, 30),
                bold=True,
                italic_guess=False,
            ),
        )

        document = Document(io.BytesIO(build_docx([region])))
        paragraph = document.paragraphs[0]
        run = paragraph.runs[0]

        self.assertEqual(paragraph.style.name, "Heading 1")
        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(run.text, "Centered heading")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.size.pt, 18)

    def test_justified_paragraph_alignment(self):
        region = Region(
            kind="text",
            bbox=[15, 46, 355, 541],
            role="paragraph",
            text="Justified body text",
            alignment="justify",
        )

        document = Document(io.BytesIO(build_docx([region])))
        paragraph = document.paragraphs[0]

        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_heading_level_uses_numbering_and_size(self):
        self.assertEqual(_infer_heading_level("I. Principles", 12), 1)
        self.assertEqual(_infer_heading_level("2.1 Scope", 12), 3)
        self.assertEqual(_infer_heading_level("Untitled", 18), 2)

    def test_embeds_figure_crop_in_docx(self):
        image = np.full((80, 120, 3), 255, dtype=np.uint8)
        image[20:60, 30:90] = (0, 0, 255)
        image_bytes = _crop_figure(image, [30, 20, 90, 60], padding_px=0)
        region = Region(kind="figure", bbox=[30, 20, 90, 60], role="figure", image_bytes=image_bytes)

        document = Document(io.BytesIO(build_docx([region])))
        self.assertEqual(len(document.inline_shapes), 1)

    def test_builds_editable_table_from_v3_markdown(self):
        region = Region(
            kind="table",
            bbox=[0, 0, 100, 100],
            role="table",
            res={"content": "| Name | Score |\n| --- | ---: |\n| Ada | 100 |"},
        )

        document = Document(io.BytesIO(build_docx([region])))
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(document.tables[0].cell(0, 0).text, "Name")
        self.assertEqual(document.tables[0].cell(1, 1).text, "100")

    def test_normalizes_v3_reading_order_and_content(self):
        class Prediction:
            json = {
                "res": {
                    "parsing_res_list": [
                        {"block_bbox": [0, 50, 100, 80], "block_label": "text", "block_content": "Second", "block_order": 1},
                        {"block_bbox": [0, 0, 100, 40], "block_label": "paragraph_title", "block_content": "# First", "block_order": 0},
                    ]
                }
            }

        class Engine:
            def predict(self, *args, **kwargs):
                return iter([Prediction()])

        previous_engine = layout._engine
        layout._engine = Engine()
        try:
            regions = layout._analyze_v3(np.zeros((100, 100, 3), dtype=np.uint8))
        finally:
            layout._engine = previous_engine

        self.assertEqual([region.order for region in regions], [0, 1])
        self.assertEqual(regions[0].kind, "paragraph_title")
        self.assertEqual(regions[1].text, "Second")


if __name__ == "__main__":
    unittest.main()
