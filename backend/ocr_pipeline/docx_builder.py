"""
Stage 7: Document assembly.
Takes the ordered list of Regions (with OCR text + style attached) and
writes a real .docx using python-docx. Tables come in as HTML from
PP-Structure's table module and get mapped to native docx tables.
"""
import io
import re
from dataclasses import dataclass
from typing import Any, Callable, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from PIL import Image

from .models import NativeTextSpan, Region
from .style import TextStyle


MAX_TABLE_ROWS = 500
MAX_TABLE_COLUMNS = 100
MAX_TABLE_CELLS = 10_000
MAX_TABLE_SPAN = 100

_HEADER_ROLES = {"page_header", "header_figure"}
_FOOTER_ROLES = {"page_footer", "page_number", "footer_figure"}

# Requiring non-whitespace next to single-dollar delimiters keeps common
# currency text such as "$5 and $6" from becoming an equation.
_INLINE_FORMULA = re.compile(
    r"(?<!\\)\$\$(?P<double>.+?)(?<!\\)\$\$"
    r"|\\\[(?P<bracket>.+?)\\\]"
    r"|\\\((?P<paren>.+?)\\\)"
    r"|(?<!\\)\$(?![\s$])(?P<single>.+?)(?<![\s\\])\$",
    re.DOTALL,
)


def _apply_run_style(run, style: TextStyle):
    run.font.size = Pt(style.font_size_pt)
    run.font.bold = style.bold
    run.font.italic = style.italic_guess
    r, g, b = style.color_rgb
    run.font.color.rgb = RGBColor(r, g, b)
    if style.font_family:
        run.font.name = style.font_family
    if style.highlight_rgb:
        _apply_highlight(run, style.highlight_rgb)


def _apply_highlight(run, color_rgb: tuple[int, int, int]) -> None:
    """Apply an exact RGB background instead of Word's limited color enum."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*color_rgb))
    run._r.get_or_add_rPr().append(shading)


def _apply_span_style(run, span: NativeTextSpan) -> None:
    run.font.size = Pt(span.font_size_pt)
    run.font.bold = span.bold
    run.font.italic = span.italic
    run.font.color.rgb = RGBColor(*span.color_rgb)
    if span.font_family:
        run.font.name = span.font_family
    if span.highlight_rgb:
        _apply_highlight(run, span.highlight_rgb)


def _add_table_from_html(doc: Document, html: str) -> bool:
    """PP-Structure table output is HTML; parse rows/cells and build a
    native docx table so it stays editable."""
    soup = BeautifulSoup(html, "html.parser")
    tables = [table for table in soup.find_all("table") if table.find_parent("table") is None]
    if len(tables) != 1 or tables[0].find("table") is not None:
        return False

    source_table = tables[0]
    rows = [row for row in source_table.find_all("tr") if row.find_parent("table") is source_table]
    if not rows or len(rows) > MAX_TABLE_ROWS:
        return False

    @dataclass
    class CellPlacement:
        row: int
        column: int
        rowspan: int
        colspan: int
        text: str
        is_header: bool

    def parse_span(cell, attribute: str) -> int | None:
        raw_value = cell.get(attribute, "1")
        if not re.fullmatch(r"[0-9]+", str(raw_value).strip()):
            return None
        value = int(raw_value)
        return value if 1 <= value <= MAX_TABLE_SPAN else None

    def cell_text(cell) -> str:
        for line_break in cell.find_all("br"):
            line_break.replace_with("\n")
        raw_text = cell.get_text("", strip=False).replace("\r\n", "\n").replace("\r", "\n")
        lines = [re.sub(r"[\t\f\v ]+", " ", line).strip() for line in raw_text.split("\n")]
        while lines and not lines[0]:
            lines.pop(0)
        while lines and not lines[-1]:
            lines.pop()
        return "\n".join(lines)

    occupied: dict[tuple[int, int], int] = {}
    placements: list[CellPlacement] = []
    for row_index, row in enumerate(rows):
        cells = row.find_all(["td", "th"], recursive=False)
        if len(cells) > MAX_TABLE_COLUMNS:
            return False
        column = 0
        for cell in cells:
            rowspan = parse_span(cell, "rowspan")
            colspan = parse_span(cell, "colspan")
            if rowspan is None or colspan is None or row_index + rowspan > len(rows):
                return False

            while column < MAX_TABLE_COLUMNS and (row_index, column) in occupied:
                column += 1
            while column < MAX_TABLE_COLUMNS and any(
                (row_index, candidate) in occupied
                for candidate in range(column, min(MAX_TABLE_COLUMNS, column + colspan))
            ):
                column += 1
            if (
                column + colspan > MAX_TABLE_COLUMNS
                or len(occupied) + rowspan * colspan > MAX_TABLE_CELLS
            ):
                return False

            placement_index = len(placements)
            for covered_row in range(row_index, row_index + rowspan):
                for covered_column in range(column, column + colspan):
                    coordinate = (covered_row, covered_column)
                    if coordinate in occupied:
                        return False
                    occupied[coordinate] = placement_index

            placements.append(CellPlacement(
                row_index,
                column,
                rowspan,
                colspan,
                cell_text(cell),
                cell.name == "th",
            ))
            column += colspan

    if not placements:
        return False

    n_cols = max(column for _, column in occupied) + 1
    if n_cols > MAX_TABLE_COLUMNS or len(rows) * n_cols > MAX_TABLE_CELLS:
        return False
    if any((row, column) not in occupied for row in range(len(rows)) for column in range(n_cols)):
        return False

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for placement in placements:
        destination = table.cell(placement.row, placement.column)
        if placement.rowspan > 1 or placement.colspan > 1:
            destination = destination.merge(
                table.cell(
                    placement.row + placement.rowspan - 1,
                    placement.column + placement.colspan - 1,
                )
            )
        _set_cell_text(destination, placement.text)
        if placement.is_header:
            _bold_cell(destination)
    if any(placement.is_header and placement.row == 0 for placement in placements):
        _mark_header_row(table.rows[0])
    return True


def _split_markdown_row(line: str) -> tuple[list[str], int]:
    """Split a Markdown row without treating escaped pipes as delimiters."""
    cells: list[str] = []
    current: list[str] = []
    delimiter_count = 0
    index = 0
    while index < len(line):
        character = line[index]
        if character == "\\" and index + 1 < len(line) and line[index + 1] in {"\\", "|"}:
            current.append(line[index + 1])
            index += 2
            continue
        if character == "|":
            cells.append("".join(current).strip())
            current = []
            delimiter_count += 1
        else:
            current.append(character)
        index += 1
    cells.append("".join(current).strip())

    stripped = line.strip()
    if stripped.startswith("|"):
        cells.pop(0)
    trailing_backslashes = len(stripped[:-1]) - len(stripped[:-1].rstrip("\\")) if stripped else 0
    if stripped.endswith("|") and trailing_backslashes % 2 == 0:
        cells.pop()
    return cells, delimiter_count


def _add_table_from_markdown(doc: Document, markdown: str) -> bool:
    """Convert a validated GitHub-style pipe table emitted by PP-StructureV3."""
    lines = [line.strip() for line in markdown.strip().splitlines()]
    if (
        len(lines) < 2
        or len(lines) - 1 > MAX_TABLE_ROWS
        or any(not line for line in lines)
    ):
        return False

    parsed_rows = [_split_markdown_row(line) for line in lines]
    if any(delimiter_count == 0 or not cells for cells, delimiter_count in parsed_rows):
        return False

    n_cols = len(parsed_rows[0][0])
    if (
        n_cols == 0
        or n_cols > MAX_TABLE_COLUMNS
        or (len(lines) - 1) * n_cols > MAX_TABLE_CELLS
        or any(len(cells) != n_cols for cells, _ in parsed_rows)
    ):
        return False

    separator = parsed_rows[1][0]
    if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in separator):
        return False
    if any(
        all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)
        for cells, _ in parsed_rows[2:]
    ):
        return False

    rows = [parsed_rows[0][0], *(cells for cells, _ in parsed_rows[2:])]
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            _set_cell_text(table.cell(row_index, column_index), value)
    for cell in table.rows[0].cells:
        _bold_cell(cell)
    _mark_header_row(table.rows[0])
    return True


def _add_table_from_rows(doc: Document, rows: object) -> bool:
    if (
        not isinstance(rows, list)
        or not rows
        or len(rows) > MAX_TABLE_ROWS
        or not all(isinstance(row, list) for row in rows)
    ):
        return False
    width = max((len(row) for row in rows), default=0)
    if (
        not width
        or width > MAX_TABLE_COLUMNS
        or len(rows) * width > MAX_TABLE_CELLS
        or any(len(row) != width for row in rows)
    ):
        return False
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            _set_cell_text(table.cell(row_index, column_index), str(value or ""))
    for cell in table.rows[0].cells:
        _bold_cell(cell)
    _mark_header_row(table.rows[0])
    return True


def _bold_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _set_cell_text(cell, text: str) -> None:
    """Set editable cell content, including any recognized inline formulas."""
    cell.text = ""
    _append_text_with_inline_math(cell.paragraphs[0], text)


def _mark_header_row(row) -> None:
    """Ask Word to repeat a recognized header row across page breaks."""
    row_properties = row._tr.get_or_add_trPr()
    if row_properties.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        row_properties.append(repeat)


def _plain_text(text: str) -> str:
    """Prevent V3 presentation markup from being emitted as literal Word text."""
    if re.search(r"</?[A-Za-z][^>]*>", text):
        return BeautifulSoup(text, "html.parser").get_text(" ", strip=True)
    return text


def _add_figure(container: Any, region: Region, section=None) -> None:
    """Embed a detected figure/logo as source pixels, scaled to the page."""
    if not region.image_bytes:
        return

    if section is None:
        section = container.sections[-1]
    max_width = section.page_width - section.left_margin - section.right_margin
    max_height = section.page_height - section.top_margin - section.bottom_margin
    with Image.open(io.BytesIO(region.image_bytes)) as image:
        # Retaining the source scale prevents a small
        # logo from becoming a page-width banner while still fitting large
        # figures within Word's printable area.
        dpi = max(72, int(region.metadata.get("source_dpi", 200)))
        source_width = Inches(image.width / dpi)
        source_height = Inches(image.height / dpi)
        preferred_scale = max(1.0, Inches(0.35) / source_width)
        fit_scale = min(max_width / source_width, max_height / source_height)
        rendered_width = max(1, int(source_width * min(preferred_scale, fit_scale)))
    paragraph = container.add_paragraph()
    paragraph.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(region.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.add_run().add_picture(
        io.BytesIO(region.image_bytes),
        # Supplying one dimension makes python-docx derive the other from the
        # source image, so the fit operation cannot distort the aspect ratio.
        width=rendered_width,
    )


def _latex_to_omml(latex: str, display: bool = False):
    """Return an OMML element, or None when a LaTeX fragment is unsupported."""
    latex = latex.strip()
    if not latex:
        return None
    try:
        import latex2mathml.converter
        import mathml2omml
        from docx.oxml import parse_xml

        mathml = latex2mathml.converter.convert(
            latex,
            display="block" if display else "inline",
        )
        omml = mathml2omml.convert(mathml)
        equation = parse_xml(omml.replace(
            "<m:oMath>",
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">',
            1,
        ))
    except Exception:
        return None
    return equation


def _strip_formula_delimiters(latex: str) -> tuple[str, bool]:
    latex = latex.strip()
    if latex.startswith("$$") and latex.endswith("$$") and len(latex) >= 4:
        return latex[2:-2].strip(), True
    if latex.startswith(r"\[") and latex.endswith(r"\]"):
        return latex[2:-2].strip(), True
    if latex.startswith("$") and latex.endswith("$") and len(latex) >= 2:
        return latex[1:-1].strip(), False
    if latex.startswith(r"\(") and latex.endswith(r"\)"):
        return latex[2:-2].strip(), False
    return latex, True


def _add_editable_formula(doc: Document, region: Region) -> bool:
    """Convert FormulaNet LaTeX to editable Word math, with caller fallback."""
    latex, display = _strip_formula_delimiters(region.formula_latex)
    equation = _latex_to_omml(latex, display=display)
    if equation is None:
        return False
    paragraph = doc.add_paragraph()
    paragraph.alignment = _paragraph_alignment(region.alignment)
    paragraph._p.append(equation)
    return True


def _append_text_with_inline_math(
    paragraph,
    text: str,
    style_run: Callable[[Any], None] | None = None,
) -> None:
    """Append mixed text/LaTeX, leaving unconvertible fragments untouched."""
    cursor = 0
    for match in _INLINE_FORMULA.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()].replace(r"\$", "$"))
            if style_run:
                style_run(run)
        latex = next(value for value in match.groupdict().values() if value is not None)
        display = match.group("double") is not None or match.group("bracket") is not None
        equation = _latex_to_omml(latex, display=display)
        if equation is None:
            run = paragraph.add_run(match.group(0))
            if style_run:
                style_run(run)
        else:
            paragraph._p.append(equation)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:].replace(r"\$", "$"))
        if style_run:
            style_run(run)


def _paragraph_alignment(value: str):
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(value, WD_ALIGN_PARAGRAPH.LEFT)


def _set_paragraph_role(doc: Document, paragraph, region: Region) -> None:
    if region.role == "document_title":
        paragraph.style = doc.styles["Title"]
    elif region.role == "heading":
        level = min(max(region.heading_level or 1, 1), 6)
        paragraph.style = doc.styles[f"Heading {level}"]
    elif region.role == "caption":
        paragraph.style = doc.styles["Caption"]
    elif region.role == "page_header":
        paragraph.style = doc.styles["Header"]
    elif region.role in {"page_footer", "page_number"}:
        paragraph.style = doc.styles["Footer"]
    elif region.role == "footnote":
        # The default python-docx template does not materialize Word's latent
        # Footnote Text style, while Caption provides the appropriate compact
        # body treatment without creating an invalid style reference.
        paragraph.style = doc.styles["Caption"]
    elif region.role == "code":
        paragraph.style = doc.styles["No Spacing"]


_LIST_PREFIX = re.compile(
    r"^\s*(?P<marker>"
    r"(?P<bullet>[\u2022\u25e6\u25aa*+-])"
    r"|(?P<number>(?:\d+(?:\.\d+)*|[A-Za-z]|[IVXLCDMivxlcdm]+)[.)]"
    r"|\((?:\d+(?:\.\d+)*|[A-Za-z]|[IVXLCDMivxlcdm]+)\))"
    r")\s+"
)


def _add_list_regions(doc: Document, region: Region, text: str) -> None:
    lines = region.lines or []
    values = [line.text for line in lines] if lines else text.splitlines()
    values = [value.strip() for value in values if value.strip()]
    if not values:
        return
    items: list[tuple[str, str]] = []
    for value in values:
        match = _LIST_PREFIX.match(value)
        cleaned = value[match.end():] if match else value
        if not match and items:
            previous, marker = items[-1]
            items[-1] = (f"{previous} {cleaned}", marker)
        else:
            items.append((cleaned, match.group("marker") if match else "\u2022"))

    for cleaned, marker in items:
        paragraph = doc.add_paragraph()
        # Word's List Number style generates its own marker and can restart or
        # change an OCR'd sequence. Keep the source marker as literal text and
        # use List Paragraph only for indentation.
        paragraph.style = doc.styles["List Paragraph"]
        paragraph.alignment = _paragraph_alignment(region.alignment)
        _append_text_with_inline_math(
            paragraph,
            f"{marker} {cleaned}",
            (lambda run: _apply_run_style(run, region.style)) if region.style else None,
        )


def _add_text_region(
    doc: Document,
    region: Region,
    text: str,
    container: Any | None = None,
) -> None:
    if region.role == "list":
        _add_list_regions(doc, region, text)
        return

    target = container if container is not None else doc
    paragraph = target.add_paragraph()
    _set_paragraph_role(doc, paragraph, region)
    paragraph.alignment = _paragraph_alignment(region.alignment)

    if region.lines:
        for line_index, line in enumerate(region.lines):
            if line_index:
                if region.role == "code":
                    paragraph.add_run().add_break()
                else:
                    # Layout OCR lines are visual wraps within one semantic
                    # block. Let Word reflow them instead of justifying every
                    # hard-broken scan line across the full text width.
                    paragraph.add_run(" ")
            for span in line.spans:
                if region.role == "code":
                    run = paragraph.add_run(span.text)
                    _apply_span_style(run, span)
                else:
                    _append_text_with_inline_math(
                        paragraph,
                        span.text,
                        lambda run, current=span: _apply_span_style(run, current),
                    )
        return

    if region.role != "code":
        # Cleanup deliberately clears OCR line objects after replacing text.
        # Any remaining newlines are visual scan wraps within this one region,
        # not authored hard breaks; Word should be allowed to reflow them.
        text = re.sub(r"\s*\n\s*", " ", text)
    style_run = (lambda run: _apply_run_style(run, region.style)) if region.style else None
    if region.role == "code":
        run = paragraph.add_run(text)
        if style_run:
            style_run(run)
    else:
        _append_text_with_inline_math(paragraph, text, style_run)


def add_regions_to_docx(doc: Document, regions: List[Region]) -> None:
    """Append one OCR page's regions to an existing document."""
    for region in regions:
        # Running page material is emitted by add_page_to_docx into the
        # section's header/footer parts, not duplicated in the body.
        if region.role in _HEADER_ROLES | _FOOTER_ROLES:
            continue
        if region.role == "table":
            content = region.res or {}
            rows = content.get("rows") if isinstance(content, dict) else None
            if rows and _add_table_from_rows(doc, rows):
                continue
            html = content.get("html", "") if isinstance(content, dict) else ""
            if not html and isinstance(content, dict):
                html = content.get("content", "")
            if html:
                if "<table" in html.lower():
                    added = _add_table_from_html(doc, html)
                else:
                    added = _add_table_from_markdown(doc, html)
                if added:
                    continue
            # A malformed or visual-only table must still survive the export.
            _add_figure(doc, region)
            continue

        if region.role == "formula":
            if not _add_editable_formula(doc, region):
                _add_figure(doc, region)
            continue

        if region.role == "figure":
            _add_figure(doc, region)
            continue

        # Algorithm/code blocks may legitimately contain HTML/XML tags and
        # LaTeX delimiters. They are source text, not Paddle presentation
        # markup, and _add_text_region deliberately emits them as literal runs.
        if region.role == "code":
            text = region.text
        else:
            text = _plain_text(region.text).strip()
        if not text.strip():
            continue

        _add_text_region(doc, region, text)


def _remove_empty_starter_paragraph(container) -> None:
    paragraphs = container.paragraphs
    if len(paragraphs) <= 1:
        return
    starter = paragraphs[0]
    if not starter.text and not starter._p.xpath(".//w:drawing"):
        starter._element.getparent().remove(starter._element)


def _add_running_regions(
    doc: Document,
    container,
    section,
    regions: List[Region],
) -> None:
    for region in regions:
        if region.role in {"header_figure", "footer_figure"}:
            _add_figure(container, region, section=section)
            continue
        text = _plain_text(region.text).strip()
        if text:
            _add_text_region(doc, region, text, container=container)
    _remove_empty_starter_paragraph(container)


def add_page_to_docx(
    doc: Document,
    regions: List[Region],
    new_section: bool = False,
) -> None:
    """Append one page and place detected running material in Word parts."""
    section = doc.add_section(WD_SECTION.NEW_PAGE) if new_section else doc.sections[-1]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    _add_running_regions(
        doc,
        section.header,
        section,
        [region for region in regions if region.role in _HEADER_ROLES],
    )
    _add_running_regions(
        doc,
        section.footer,
        section,
        [region for region in regions if region.role in _FOOTER_ROLES],
    )
    add_regions_to_docx(doc, regions)


def build_docx(regions: List[Region]) -> bytes:
    """Build a DOCX from one OCR page's ordered regions."""
    doc = Document()
    add_page_to_docx(doc, regions)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
