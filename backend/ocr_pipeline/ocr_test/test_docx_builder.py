import io
import unittest
from zipfile import ZipFile

import cv2
import numpy as np

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ocr_pipeline.docx_builder import build_docx
from ocr_pipeline import layout
from ocr_pipeline.layout import Region
from ocr_pipeline.models import DocumentIR, NativeTextLine, NativeTextSpan, PageIR
from ocr_pipeline.pipeline import (
    _assign_heading_hierarchy,
    _classify_region,
    _crop_figure,
    _infer_alignment,
    _infer_heading_level,
    document_ir_to_docx,
)
from ocr_pipeline.style import TextStyle


class DocxBuilderTests(unittest.TestCase):
    def test_classifies_layout_regions_and_alignment(self):
        self.assertEqual(_classify_region(Region(kind="title", bbox=[0, 0, 1, 1])), "heading")
        self.assertEqual(_classify_region(Region(kind="text", bbox=[0, 0, 1, 1])), "paragraph")
        self.assertEqual(_classify_region(Region(kind="list", bbox=[0, 0, 1, 1])), "list")
        self.assertEqual(_classify_region(Region(kind="image", bbox=[0, 0, 1, 1])), "figure")
        self.assertEqual(_classify_region(Region(kind="formula", bbox=[0, 0, 1, 1])), "formula")
        self.assertEqual(_classify_region(Region(kind="doc_title", bbox=[0, 0, 1, 1])), "document_title")
        self.assertEqual(_classify_region(Region(kind="figure_title", bbox=[0, 0, 1, 1])), "caption")
        self.assertEqual(_classify_region(Region(kind="header", bbox=[0, 0, 1, 1])), "page_header")
        self.assertEqual(_classify_region(Region(kind="header_image", bbox=[0, 0, 1, 1])), "header_figure")
        self.assertEqual(_classify_region(Region(kind="footer_image", bbox=[0, 0, 1, 1])), "footer_figure")
        self.assertEqual(_classify_region(Region(kind="formula_number", bbox=[0, 0, 1, 1])), "caption")
        self.assertEqual(_infer_alignment([420, 80, 1180, 160], 1600), "center")
        self.assertEqual(_infer_alignment([120, 180, 1320, 230], 1600), "left")
        # Full-width region with narrow, near-symmetric margins on both
        # sides -- conventionally justified body text (verified against
        # PaddleOCR's own native docx export on an equivalent block).
        self.assertEqual(_infer_alignment([15, 46, 355, 541], 370), "justify")

    def test_preserves_region_semantics_and_text_style(self):
        region = Region(
            kind="title",
            bbox=[300, 80, 1300, 160],
            role="heading",
            text="Centered heading",
            alignment="center",
            style=TextStyle(
                font_size_pt=18,
                color_rgb=(10, 20, 30),
                bold=True,
                italic_guess=False,
            ),
        )

        document = Document(io.BytesIO(build_docx([region])))
        paragraph = document.paragraphs[0]
        run = paragraph.runs[0]

        self.assertEqual(paragraph.style.name, "Heading 1")
        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(run.text, "Centered heading")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.size.pt, 18)

    def test_justified_paragraph_alignment(self):
        region = Region(
            kind="text",
            bbox=[15, 46, 355, 541],
            role="paragraph",
            text="Justified body text",
            alignment="justify",
        )

        document = Document(io.BytesIO(build_docx([region])))
        paragraph = document.paragraphs[0]

        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_reflows_visual_newlines_when_cleanup_removed_line_geometry(self):
        region = Region(
            kind="text",
            bbox=[0, 0, 100, 30],
            role="paragraph",
            text="A visual OCR\nwrap should reflow.",
        )
        result = build_docx([region])
        document = Document(io.BytesIO(result))
        self.assertEqual(document.paragraphs[0].text, "A visual OCR wrap should reflow.")
        with ZipFile(io.BytesIO(result)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertNotIn("<w:br", document_xml)

    def test_heading_level_uses_numbering_and_size(self):
        self.assertEqual(_infer_heading_level("I. Principles", 12), 1)
        self.assertEqual(_infer_heading_level("2.1 Scope", 12), 2)
        self.assertEqual(_infer_heading_level("2.1.3 Limits", 12), 3)
        self.assertEqual(_infer_heading_level("Untitled", 18), 2)

    def test_heading_levels_are_resolved_across_the_document(self):
        regions = [
            Region(kind="paragraph_title", bbox=[0, 0, 10, 10], role="heading", text="Overview", style=TextStyle(20, (0, 0, 0), True, False)),
            Region(kind="paragraph_title", bbox=[0, 20, 10, 30], role="heading", text="Details", style=TextStyle(14, (0, 0, 0), True, False)),
            Region(kind="paragraph_title", bbox=[0, 40, 10, 50], role="heading", text="3.2 Explicit", style=TextStyle(20, (0, 0, 0), True, False)),
        ]
        _assign_heading_hierarchy(regions)
        self.assertEqual([region.heading_level for region in regions], [1, 2, 2])
        self.assertEqual(regions[2].metadata["heading_level_source"], "numbering")

    def test_embeds_figure_crop_in_docx(self):
        image = np.full((80, 120, 3), 255, dtype=np.uint8)
        image[20:60, 30:90] = (0, 0, 255)
        image_bytes = _crop_figure(image, [30, 20, 90, 60], padding_px=0)
        region = Region(kind="figure", bbox=[30, 20, 90, 60], role="figure", image_bytes=image_bytes)

        document = Document(io.BytesIO(build_docx([region])))
        self.assertEqual(len(document.inline_shapes), 1)

    def test_constrains_tall_figure_to_printable_height_without_distortion(self):
        image = np.full((1200, 100, 3), 255, dtype=np.uint8)
        success, encoded = cv2.imencode(".png", image)
        self.assertTrue(success)
        region = Region(
            kind="figure",
            bbox=[0, 0, 100, 1200],
            role="figure",
            image_bytes=encoded.tobytes(),
            metadata={"source_dpi": 100},
        )

        document = Document(io.BytesIO(build_docx([region])))
        shape = document.inline_shapes[0]
        section = document.sections[0]
        printable_height = section.page_height - section.top_margin - section.bottom_margin

        self.assertLessEqual(shape.height, printable_height)
        self.assertAlmostEqual(shape.width / shape.height, 100 / 1200, places=3)

    def test_constrains_wide_figure_to_printable_width_without_distortion(self):
        image = np.full((100, 1200, 3), 255, dtype=np.uint8)
        success, encoded = cv2.imencode(".png", image)
        self.assertTrue(success)
        region = Region(
            kind="figure",
            bbox=[0, 0, 1200, 100],
            role="figure",
            image_bytes=encoded.tobytes(),
            metadata={"source_dpi": 100},
        )

        document = Document(io.BytesIO(build_docx([region])))
        shape = document.inline_shapes[0]
        section = document.sections[0]
        printable_width = section.page_width - section.left_margin - section.right_margin

        self.assertLessEqual(shape.width, printable_width)
        self.assertAlmostEqual(shape.width / shape.height, 1200 / 100, places=3)

    def test_builds_editable_table_from_v3_markdown(self):
        region = Region(
            kind="table",
            bbox=[0, 0, 100, 100],
            role="table",
            res={"content": "| Name | Score |\n| --- | ---: |\n| Ada | 100 |"},
        )

        document = Document(io.BytesIO(build_docx([region])))
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(document.tables[0].cell(0, 0).text, "Name")
        self.assertEqual(document.tables[0].cell(1, 1).text, "100")

    def test_preserves_formula_and_unparsed_table_as_images(self):
        image = np.full((80, 120, 3), 255, dtype=np.uint8)
        image[20:60, 30:90] = (0, 0, 255)
        crop = _crop_figure(image, [30, 20, 90, 60], padding_px=0)
        regions = [
            Region(kind="formula", bbox=[30, 20, 90, 60], role="formula", image_bytes=crop),
            Region(kind="table", bbox=[30, 20, 90, 60], role="table", res={"content": "not a table"}, image_bytes=crop),
        ]

        document = Document(io.BytesIO(build_docx(regions)))
        self.assertEqual(len(document.tables), 0)
        self.assertEqual(len(document.inline_shapes), 2)

    def test_emits_formula_latex_as_editable_word_math(self):
        region = Region(
            kind="formula",
            bbox=[0, 0, 100, 30],
            role="formula",
            formula_latex=r"x^2 + y^2 = z^2",
        )
        result = build_docx([region])
        with ZipFile(io.BytesIO(result)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("<m:oMath", document_xml)

    def test_honors_standalone_formula_alignment(self):
        regions = [
            Region(
                kind="formula",
                bbox=[0, 0, 100, 30],
                role="formula",
                formula_latex="x=1",
                alignment="left",
            ),
            Region(
                kind="formula",
                bbox=[0, 40, 100, 70],
                role="formula",
                formula_latex="y=2",
                alignment="right",
            ),
        ]
        document = Document(io.BytesIO(build_docx(regions)))

        self.assertEqual(document.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(document.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.RIGHT)

    def test_emits_inline_latex_without_converting_currency(self):
        region = Region(
            kind="text",
            bbox=[0, 0, 100, 30],
            role="paragraph",
            text=r"Energy $E=mc^2$ and \(x+1\); prices are $5 and $6.",
        )
        result = build_docx([region])
        with ZipFile(io.BytesIO(result)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertEqual(document_xml.count("<m:oMath>"), 2)
        self.assertIn("$5 and $6.", Document(io.BytesIO(result)).paragraphs[0].text)

    def test_emits_inline_latex_in_lists_but_not_code(self):
        regions = [
            Region(kind="list", bbox=[0, 0, 100, 30], role="list", text=r"- Value $x^2$"),
            Region(kind="algorithm", bbox=[0, 40, 100, 70], role="code", text=r"price = '$x^2$'"),
        ]
        result = build_docx(regions)
        with ZipFile(io.BytesIO(result)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        document = Document(io.BytesIO(result))

        self.assertEqual(document_xml.count("<m:oMath>"), 1)
        self.assertEqual(document.paragraphs[0].text, "- Value ")
        self.assertEqual(document.paragraphs[1].text, r"price = '$x^2$'")

    def test_preserves_numbered_and_alphabetic_list_markers_literally(self):
        region = Region(
            kind="list",
            bbox=[0, 0, 100, 80],
            role="list",
            text="3) Third\n7. Seventh\nB) Alpha\n(c) Lower",
        )
        result = build_docx([region])
        document = Document(io.BytesIO(result))
        with ZipFile(io.BytesIO(result)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs],
            ["3) Third", "7. Seventh", "B) Alpha", "(c) Lower"],
        )
        self.assertTrue(all(paragraph.style.name == "List Paragraph" for paragraph in document.paragraphs))
        self.assertNotIn("<w:numPr>", document_xml)

    def test_keeps_html_and_latex_literal_in_code_regions(self):
        region = Region(
            kind="algorithm",
            bbox=[0, 0, 100, 30],
            role="code",
            text="  <div class='formula'>$x^2$</div>\n",
        )
        result = build_docx([region])
        document = Document(io.BytesIO(result))
        with ZipFile(io.BytesIO(result)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertEqual(document.paragraphs[0].text, "  <div class='formula'>$x^2$</div>\n")
        self.assertNotIn("<m:oMath>", document_xml)

    def test_places_running_text_in_real_section_headers_and_footers(self):
        pages = DocumentIR(pages=[
            PageIR(page_index=0, regions=[
                Region(kind="header", bbox=[0, 0, 10, 10], role="page_header", text="First header"),
                Region(kind="text", bbox=[0, 20, 10, 30], role="paragraph", text="First body"),
                Region(kind="footer", bbox=[0, 90, 10, 100], role="page_footer", text="First footer"),
            ]),
            PageIR(page_index=1, regions=[
                Region(kind="header", bbox=[0, 0, 10, 10], role="page_header", text="Second header"),
                Region(kind="text", bbox=[0, 20, 10, 30], role="paragraph", text="Second body"),
                Region(kind="number", bbox=[0, 90, 10, 100], role="page_number", text="2"),
            ]),
        ])
        document = Document(io.BytesIO(document_ir_to_docx(pages)))

        self.assertEqual(len(document.sections), 2)
        self.assertEqual(document.sections[0].header.paragraphs[0].text, "First header")
        self.assertEqual(document.sections[1].header.paragraphs[0].text, "Second header")
        self.assertEqual(document.sections[0].footer.paragraphs[0].text, "First footer")
        self.assertEqual(document.sections[1].footer.paragraphs[0].text, "2")
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["First body", "Second body"],
        )

    def test_places_detected_header_image_in_header_part(self):
        image = np.full((30, 80, 3), 255, dtype=np.uint8)
        crop = _crop_figure(image, [0, 0, 80, 30], padding_px=0)
        region = Region(
            kind="header_image",
            bbox=[0, 0, 80, 30],
            role="header_figure",
            image_bytes=crop,
        )
        document = Document(io.BytesIO(build_docx([region])))

        self.assertTrue(document.sections[0].header._element.xpath(".//w:drawing"))
        self.assertEqual(len(document.inline_shapes), 0)

    def test_strips_html_presentation_markup_from_text(self):
        region = Region(
            kind="text",
            bbox=[0, 0, 10, 10],
            role="paragraph",
            text='<div><img src="formula.jpg" alt="formula" /></div>',
        )
        document = Document(io.BytesIO(build_docx([region])))
        self.assertEqual(document.paragraphs, [])

    def test_normalizes_v3_reading_order_and_content(self):
        class Prediction:
            json = {
                "res": {
                    "layout_det_res": {
                        "boxes": [
                            {"label": "paragraph_title", "score": 0.91, "coordinate": [0, 0, 100, 40]},
                            {"label": "text", "score": 0.82, "coordinate": [0, 50, 100, 80]},
                        ]
                    },
                    "overall_ocr_res": {
                        "rec_texts": ["First", "Second"],
                        "rec_scores": [0.93, 0.88],
                        "rec_boxes": [[5, 5, 90, 30], [5, 55, 90, 75]],
                    },
                    "parsing_res_list": [
                        {"block_bbox": [0, 50, 100, 80], "block_label": "text", "block_content": "Second", "block_order": 1},
                        {"block_bbox": [0, 0, 100, 40], "block_label": "paragraph_title", "block_content": "# First", "block_order": 0},
                    ]
                }
            }

        class Engine:
            def predict(self, *args, **kwargs):
                return iter([Prediction()])

        previous_engine = layout._engine
        layout._engine = Engine()
        try:
            regions = layout._analyze_v3(np.zeros((100, 100, 3), dtype=np.uint8))
        finally:
            layout._engine = previous_engine

        self.assertEqual([region.order for region in regions], [0, 1])
        self.assertEqual(regions[0].kind, "paragraph_title")
        self.assertAlmostEqual(regions[0].confidence, 0.91)
        self.assertEqual(regions[0].metadata["paddle_ocr_lines"][0]["text"], "First")
        self.assertEqual(regions[1].text, "Second")

    def test_preserves_mixed_native_pdf_runs(self):
        region = Region(
            kind="text",
            bbox=[0, 0, 100, 20],
            role="paragraph",
            text="Normal bold",
            native_text=True,
            lines=[NativeTextLine(
                bbox=[0, 0, 100, 20],
                block_index=0,
                spans=[
                    NativeTextSpan("Normal ", [0, 0, 50, 20], 10, font_family="Helvetica"),
                    NativeTextSpan("bold", [50, 0, 100, 20], 12, font_family="Times", bold=True, color_rgb=(20, 30, 40)),
                ],
            )],
        )
        document = Document(io.BytesIO(build_docx([region])))
        runs = document.paragraphs[0].runs
        self.assertEqual([run.text for run in runs], ["Normal ", "bold"])
        self.assertFalse(runs[0].bold)
        self.assertTrue(runs[1].bold)
        self.assertEqual(runs[1].font.size.pt, 12)


if __name__ == "__main__":
    unittest.main()
