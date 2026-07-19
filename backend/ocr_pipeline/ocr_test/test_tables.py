import io
import unittest
from zipfile import ZipFile

from docx import Document
from PIL import Image

from ocr_pipeline.docx_builder import (
    _add_table_from_html,
    _add_table_from_markdown,
    _add_table_from_rows,
    build_docx,
)
from ocr_pipeline.models import Region


class HtmlTableTests(unittest.TestCase):
    def test_builds_rowspan_colspan_and_line_breaks(self):
        html = """
            <table>
              <tr>
                <th rowspan="2"><strong>Team</strong><br>Name</th>
                <th colspan="2">Scores</th>
              </tr>
              <tr><th>Q1</th><th>Q2</th></tr>
              <tr><td>Ada</td><td>10</td><td>20</td></tr>
            </table>
        """
        document = Document()

        self.assertTrue(_add_table_from_html(document, html))
        self.assertEqual(len(document.tables), 1)

        table = document.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 3)
        self.assertEqual(table.cell(0, 0).text, "Team\nName")
        self.assertEqual(table.cell(1, 0).text, "Team\nName")
        self.assertEqual(table.cell(0, 1).text, "Scores")
        self.assertEqual(table.cell(0, 2).text, "Scores")
        self.assertEqual(table.cell(2, 2).text, "20")
        self.assertTrue(table.cell(0, 0).paragraphs[0].runs[0].bold)

    def test_rejects_invalid_or_non_rectangular_html_before_adding_table(self):
        malformed_tables = [
            "<table><tr><td rowspan='3'>A</td></tr><tr><td>B</td></tr></table>",
            "<table><tr><td>A</td><td>B</td></tr><tr><td>C</td></tr></table>",
            "<table><tr><td colspan='nope'>A</td></tr></table>",
            "<table></table>",
            "<table><tr><td>A<table><tr><td>B</td></tr></table></td></tr></table>",
        ]

        for html in malformed_tables:
            with self.subTest(html=html):
                document = Document()
                self.assertFalse(_add_table_from_html(document, html))
                self.assertEqual(len(document.tables), 0)

    def test_rejects_excessive_html_spans_before_allocating_a_table(self):
        document = Document()
        self.assertFalse(_add_table_from_html(
            document,
            "<table><tr><td colspan='101'>oversized</td></tr></table>",
        ))
        self.assertEqual(len(document.tables), 0)


class MarkdownTableTests(unittest.TestCase):
    def test_builds_native_pdf_table_rows(self):
        region = Region(
            kind="table",
            bbox=[0, 0, 40, 20],
            role="table",
            res={"rows": [["Name", "Score"], ["Ada", "100"]]},
        )
        document = Document(io.BytesIO(build_docx([region])))
        self.assertEqual(document.tables[0].cell(1, 0).text, "Ada")
        self.assertEqual(document.tables[0].cell(1, 1).text, "100")
        self.assertTrue(document.tables[0].cell(0, 0).paragraphs[0].runs[0].bold)

    def test_supports_escaped_pipes(self):
        markdown = (
            r"| Expression | Meaning |" "\n"
            r"| :--- | ---: |" "\n"
            r"| a \| b | logical choice |"
        )
        document = Document()

        self.assertTrue(_add_table_from_markdown(document, markdown))
        table = document.tables[0]
        self.assertEqual(table.cell(0, 0).text, "Expression")
        self.assertEqual(table.cell(1, 0).text, "a | b")
        self.assertEqual(table.cell(1, 1).text, "logical choice")

    def test_converts_formula_inside_table_cell_to_editable_math(self):
        document = Document()
        self.assertTrue(_add_table_from_markdown(
            document,
            "| Expression | Meaning |\n| --- | --- |\n| $x^2$ | square |",
        ))
        buffer = io.BytesIO()
        document.save(buffer)
        with ZipFile(io.BytesIO(buffer.getvalue())) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("<m:oMath>", document_xml)

    def test_rejects_malformed_markdown_before_adding_table(self):
        malformed_tables = [
            "| A | B |\n| one | two |",
            "| A | B |\n| --- | --- |\n| only one |",
            "| A | B |\n| -- | --- |\n| 1 | 2 |",
            "| A | B |\n| --- | --- |\n\n| 1 | 2 |",
            "ordinary text with | one pipe",
        ]

        for markdown in malformed_tables:
            with self.subTest(markdown=markdown):
                document = Document()
                self.assertFalse(_add_table_from_markdown(document, markdown))
                self.assertEqual(len(document.tables), 0)

    def test_rejects_excessive_structured_table_dimensions(self):
        document = Document()
        oversized_rows = [[str(column) for column in range(101)]]
        self.assertFalse(_add_table_from_rows(document, oversized_rows))
        self.assertEqual(len(document.tables), 0)

    def test_malformed_table_uses_image_fallback(self):
        image_buffer = io.BytesIO()
        Image.new("RGB", (40, 20), "white").save(image_buffer, format="PNG")
        region = Region(
            kind="table",
            bbox=[0, 0, 40, 20],
            role="table",
            res={"content": "| A | B |\n| -- | --- |\n| 1 | 2 |"},
            image_bytes=image_buffer.getvalue(),
        )

        document = Document(io.BytesIO(build_docx([region])))

        self.assertEqual(len(document.tables), 0)
        self.assertEqual(len(document.inline_shapes), 1)


if __name__ == "__main__":
    unittest.main()
