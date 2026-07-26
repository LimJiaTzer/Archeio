from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn


INPUT = "/Users/limjiatzer/Documents/CS/Projects/Archeio/document-edits/Archeío_Readme_updated.docx"


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


document = Document(INPUT)
paragraphs = document.paragraphs
start = next(index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "Frontend")
end = next(index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "Milestone Plans")

replacement = [
    ("Frontend", "Heading 2"),
    ("React 19 + React DOM", "normal"),
    ("Vite (build tool)", "normal"),
    ("React Router DOM", "normal"),
    ("Tailwind CSS + PostCSS", "normal"),
    ("Framer Motion (animations)", "normal"),
    ("Lucide React (icons)", "normal"),
    ("clsx + tailwind-merge (utility libraries)", "normal"),
    ("Vercel (for deployment)", "normal"),
    ("Media Processing (client-side, WebAssembly)", "Heading 2"),
    ("@ffmpeg/ffmpeg + @ffmpeg/core + @ffmpeg/util — FFmpeg compiled to WebAssembly, running entirely in the browser for video and audio conversion", "normal"),
    ("HTML5 Canvas API — image format conversion (the canvas-in-Vite approach you mentioned)", "normal"),
    ("Document Processing (client-side)", "Heading 2"),
    ("jspdf — PDF generation", "normal"),
    ("pdf-lib — PDF manipulation", "normal"),
    ("pdfjs-dist — PDF reading and rendering", "normal"),
    ("html2pdf.js — HTML to PDF conversion", "normal"),
    ("mammoth — DOCX to HTML/plain text conversion", "normal"),
    ("Backend [Planned]", "Heading 2"),
    ("Python (FastAPI / Flask)", "normal"),
    ("React Native", "normal"),
    ("Oracle VM Free", "normal"),
    ("", "Heading 1"),
]

for paragraph, (text, style) in zip(paragraphs[start:end], replacement):
    paragraph.text = text
    paragraph.style = style

# Keep the supplied entries as one consistent bulleted list per subsection.
# Clearing inherited numbering first prevents former content roles from leaving
# bullets on headings or removing bullets from newly inserted entries.
bullet_num_pr = deepcopy(paragraphs[start + 1]._p.pPr.find(qn("w:numPr")))
for paragraph, (text, style) in zip(paragraphs[start:end], replacement):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    old_num_pr = paragraph_properties.find(qn("w:numPr"))
    if old_num_pr is not None:
        paragraph_properties.remove(old_num_pr)
    if style == "normal" and text:
        paragraph_properties.append(deepcopy(bullet_num_pr))

# The original section has one now-unused spacer paragraph; remove it so the
# following Milestone Plans heading retains its original separation.
unused_spacer = document.paragraphs[start + len(replacement)]
if not unused_spacer.text.strip():
    remove_paragraph(unused_spacer)
document.save(INPUT)
